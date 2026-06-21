"""Generate PROJECT_DOCUMENTATION.docx — a full, in-depth project report.

Run:  python docs/generate_project_doc.py
Output: docs/PROJECT_DOCUMENTATION.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent / "PROJECT_DOCUMENTATION.docx"

doc = Document()

# ---- base styles -------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def p(text, bold=False, italic=False):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    return par


def bullet(text, bold_prefix=None):
    par = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = par.add_run(bold_prefix)
        r.bold = True
        par.add_run(text)
    else:
        par.add_run(text)
    return par


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def code(text):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return par


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hcell in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(hcell)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


# ============================== TITLE PAGE ====================================
title = doc.add_heading("Face Detection + Recognition with Signal Visualization", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Full Project Documentation\nHaar + LBPH  vs  DNN-SSD + OpenFace Embeddings")
r.italic = True
r.font.size = Pt(13)
doc.add_paragraph()

p("This document explains the project end to end: what it does, how every piece works, the "
  "signal-processing theory behind it, the file/module layout, how to install and run it, the "
  "design decisions, and likely questions with answers. It is written so you can present and "
  "defend the project in depth without missing any point.")

doc.add_page_break()

# table of contents (manual)
h1("Contents")
for i, sec in enumerate([
    "1. Project Overview",
    "2. Detection vs Recognition (core concepts)",
    "3. Pipeline A — Haar Cascade + LBPH",
    "4. Pipeline B — DNN-SSD + OpenFace Embeddings",
    "5. Signal Visualization 1 — RGB to Grayscale (Haar path)",
    "6. Signal Visualization 2 — The BLOB (DNN path)",
    "7. Enrollment (registering people)",
    "8. The Comparison — which pipeline is best",
    "9. Project Structure (files & folders)",
    "10. Module-by-Module Breakdown",
    "11. The 5-Tab User Interface",
    "12. Installation & Running",
    "13. Data Flow (end to end)",
    "14. Design Decisions & Why",
    "15. Key Technical Notes & Gotchas",
    "16. Limitations & Honest Caveats",
    "17. Possible Future Improvements",
    "18. Verification Performed",
    "19. Glossary",
    "20. One-Paragraph Summary (for oral explanation)",
], 1):
    doc.add_paragraph(sec, style="List Bullet")

doc.add_page_break()

# ============================== 1. OVERVIEW ===================================
h1("1. Project Overview")
p("This project began as two simple OpenCV scripts that only DETECTED faces (drew rectangles "
  "around them). It was expanded into a complete Streamlit web application that performs four "
  "distinct jobs:")
numbered("Detection — finds WHERE faces are in a webcam frame.")
numbered("Recognition — identifies WHO the face is (by name), not just where.")
numbered("Signal visualization — shows the signal-processing internals: how the RGB camera "
         "signal becomes grayscale (Haar path), and how the SSD network's BLOB input signal "
         "varies between different faces (DNN path).")
numbered("Comparison — runs two complete pipelines side by side and declares which is best.")
p("Two complete pipelines are compared:")
table(
    ["Pipeline", "Detector", "Recognizer", "Box colour"],
    [
        ["A", "Haar Cascade (Viola-Jones)", "LBPH (Local Binary Patterns Histograms)", "Blue"],
        ["B", "DNN-SSD (res10 Caffe model)", "OpenFace 128-D embeddings", "Green"],
    ],
)
p("Key design decision: everything stays within the OpenCV family — no dlib or face_recognition "
  "library is required. This keeps installation light and avoids heavy native compilation.", bold=True)

# ============================== 2. CONCEPTS ===================================
h1("2. Detection vs Recognition (core concepts)")
bullet("answers 'Is there a face, and where?' -> outputs a bounding box.", bold_prefix="Detection ")
bullet("answers 'Whose face is this?' -> outputs a name (or 'Unknown').", bold_prefix="Recognition ")
p("Recognition needs REFERENCE DATA. You must first ENROLL known people (store sample images of "
  "their faces), then TRAIN a recognizer on them. At runtime, every detected face is compared "
  "against the enrolled people, and either matched to a name or rejected as 'Unknown'.")

# ============================== 3. PIPELINE A =================================
h1("3. Pipeline A — Haar Cascade + LBPH")
h2("3.1 Haar Cascade Detection (Viola-Jones algorithm)")
bullet("Operates on the GRAYSCALE version of the image.")
bullet("Uses 'Haar-like features' — patterns of adjacent light/dark rectangles — that respond to "
       "facial structure (e.g. eyes are darker than cheeks).")
bullet("Features are evaluated extremely fast using an INTEGRAL IMAGE (a precomputed sum table).")
bullet("A CASCADE of classifiers rejects clearly non-face regions in early, cheap stages and only "
       "spends heavy computation on promising regions.")
bullet("Parameters used in this project: scaleFactor=1.1, minNeighbors=5, minSize=(30,30).")
bullet("Fast and lightweight, but weaker on tilted faces, side profiles, and poor lighting.")
bullet("Uses OpenCV's bundled haarcascade_frontalface_default.xml — NO external model file needed.")
h2("3.2 LBPH Recognition (Local Binary Patterns Histograms)")
bullet("Pairs naturally with Haar because it is also HISTOGRAM-based, which fits this project's "
       "'signals' theme.")
bullet("For each pixel, LBP compares it to its surrounding neighbours and forms a binary pattern; "
       "these patterns are summarised into HISTOGRAMS over image regions.")
bullet("Training builds a histogram model per enrolled person; prediction finds the closest match "
       "by comparing histograms.")
bullet("Returns a DISTANCE score: LOWER = MORE confident. If distance exceeds the threshold "
       "(70.0), the face is labelled 'Unknown'.")
bullet("Critical install fact: LBPH lives in cv2.face, which ships ONLY in opencv-contrib-python "
       "(not plain opencv-python).")

# ============================== 4. PIPELINE B =================================
h1("4. Pipeline B — DNN-SSD + OpenFace Embeddings")
h2("4.1 DNN-SSD Detection (res10 300x300 Caffe model)")
bullet("A deep neural network (Single Shot Detector) trained specifically for face detection.")
bullet("The frame is converted into a BLOB (see Section 6) — resized to 300x300 with channel "
       "means subtracted — and passed through the network in one forward pass.")
bullet("Outputs detections, each with a CONFIDENCE score; only detections above 0.5 are kept.")
bullet("More accurate than Haar across angles, lighting, and distance, but heavier to compute.")
bullet("Requires two files: deploy.prototxt (network architecture) and "
       "res10_300x300_ssd_iter_140000.caffemodel (trained weights).")
h2("4.2 OpenFace Embedding Recognition")
bullet("The cropped face is passed through the OpenFace net (openface_nn4.small2.v1.t7), producing "
       "a 128-DIMENSIONAL vector ('embedding') that numerically represents the face.")
bullet("During enrollment, each person gets an AVERAGE embedding (a 'centroid') computed from all "
       "their sample images.")
bullet("At runtime, the new face's embedding is compared to every centroid using COSINE DISTANCE "
       "(1 - cosine similarity). The nearest centroid wins.")
bullet("Returns a distance: LOWER = MORE confident. Above threshold (0.45) -> 'Unknown'.")
bullet("Much more robust to lighting and pose than LBPH, but needs the embedding model file.")

# ============================== 5. RGB -> GRAY ================================
h1("5. Signal Visualization 1 — RGB to Grayscale (Haar path)")
p("The Haar detector works on a single grayscale signal, so this view shows how three colour "
  "signals collapse into one. The luma formula (ITU-R BT.601) is:")
code("Gray = 0.299 x R + 0.587 x G + 0.114 x B")
p("Green is weighted most heavily because the human eye is most sensitive to green light.")
p("The 'Signal Lab' tab displays:")
bullet("The original image (RGB).")
bullet("The R, G, and B channels shown separately as individual signals.")
bullet("The grayscale result plus the exact weights applied.")
bullet("Intensity histograms (256 bins) for R, G, B and Gray — showing the distribution of pixel "
       "intensities in each signal.")
p("Input can be a webcam snapshot or an uploaded image.")

# ============================== 6. BLOB ======================================
h1("6. Signal Visualization 2 — The BLOB (DNN path)")
p("Central insight: the DNN never sees your raw camera frame. It sees a preprocessed tensor "
  "called a BLOB. This is how the blob is built:", bold=True)
code('blob = cv2.dnn.blobFromImage(\n'
     '    cv2.resize(frame, (300, 300)),   # resize to network input size\n'
     '    1.0,                             # scale factor\n'
     '    (300, 300),                      # spatial size\n'
     '    (104.0, 177.0, 123.0)            # mean subtraction (B, G, R)\n'
     ')')
bullet("Resulting shape: (1, 3, 300, 300) = (batch, channels, height, width).")
bullet("Each channel has a fixed MEAN subtracted (B-104, G-177, R-123) to normalise lighting and "
       "match how the network was trained.")
bullet("These three mean-subtracted channels ARE the signal the network processes — and they "
       "change with every different face and lighting condition.")
p("The 'BLOB Lab' tab displays:")
bullet("The blob tensor SHAPE.")
bullet("The resized 300x300 image (pre-subtraction, for reference).")
bullet("Three channel HEATMAPS (the mean-subtracted B, G, R channels, colour-mapped).")
bullet("A per-channel STATISTICS table (min, max, mean, std) — these numbers visibly shift when a "
       "different face is presented, which is exactly what demonstrates 'how the BLOB varies with "
       "different signals'.")
bullet("Value-distribution HISTOGRAMS for each channel.")
p("How to demo it: capture one face and note the stats; capture a different face (or change "
  "lighting) and compare — the means and standard deviations move. That is the proof of signal "
  "variation.", italic=True)

# ============================== 7. ENROLLMENT ================================
h1("7. Enrollment (registering people)")
p("Recognition requires reference faces. Two enrollment methods are supported:")
numbered("Webcam capture — type a name, capture N samples (default 20) live; the DNN detector "
         "auto-crops the largest face from each frame and saves it.")
numbered("Image folder import — point at a folder of existing images of a person; they are copied "
         "into the dataset.")
p("All faces are stored under dataset/<name>/*.jpg. After enrolling, clicking TRAIN will:")
bullet("Read every enrolled face.")
bullet("Rebuild BOTH recognizers (the LBPH model and the embedding centroids).")
bullet("Save them to artifacts/ (lbph.yml, label_map.json, embeddings.npz).")

# ============================== 8. COMPARISON ===============================
h1("8. The Comparison — which pipeline is best")
p("The Compare tab evaluates both pipelines on three criteria:")
numbered("Recognition accuracy — train on a subset of enrolled faces, test on the held-out "
         "remainder, measure how often the correct person is identified.")
numbered("Detection robustness — average number of faces detected per dataset image.")
numbered("Speed (FPS) — average frames-per-second for detect + identify.")
p("A winner is chosen PER criterion, and the OVERALL winner is decided by majority vote across "
  "the three.")
p("Honesty caveat (important for the demo/viva): accuracy is measured on a held-out split of the "
  "ENROLLED faces, not an independent labelled test set — so it is INDICATIVE, not a formal "
  "benchmark. The UI states this explicitly. To make it fairer, enroll several people with a "
  "handful of varied images each.", bold=True)

# ============================== 9. STRUCTURE ================================
h1("9. Project Structure (files & folders)")
code(
    "Face-detection-using-DNN-model-main/\n"
    "|- app.py                       # Streamlit entry point - 5-tab router\n"
    "|- requirements.txt             # opencv-contrib-python, numpy, streamlit\n"
    "|- .gitignore                   # ignores models/, dataset/, artifacts/\n"
    "|- face/                        # core logic (no UI, reusable, testable)\n"
    "|  |- config.py                 # all paths + thresholds, anchored to root\n"
    "|  |- detectors.py              # HaarDetector, DnnSsdDetector\n"
    "|  |- signals.py                # RGB->gray views + BLOB unpacking\n"
    "|  |- recognizers.py            # LbphRecognizer + EmbeddingRecognizer\n"
    "|  |- enroll.py                 # webcam capture + folder import + train_all()\n"
    "|  '- compare.py                # accuracy / robustness / FPS -> winner\n"
    "|- ui/                          # one render() per Streamlit tab\n"
    "|  |- common.py                 # BGR->RGB, Camera class, cached singletons\n"
    "|  |- live.py                   # Tab 1: live detect + recognize\n"
    "|  |- signal_lab.py             # Tab 2: RGB -> grayscale visualization\n"
    "|  |- blob_lab.py               # Tab 3: BLOB variation visualization\n"
    "|  |- enroll_ui.py              # Tab 4: enroll + train\n"
    "|  '- compare_ui.py             # Tab 5: comparison results\n"
    "|- scripts/download_models.py   # fetches the 3 model files into models/\n"
    "|- docs/                        # this documentation + generator\n"
    "|- models/      (gitignored)    # downloaded model weights\n"
    "|- dataset/     (gitignored)    # enrolled faces, per person\n"
    "|- artifacts/   (gitignored)    # trained recognizer files\n"
    "|- Face_Detection_Project(PBL).py        # legacy Haar-only demo (kept)\n"
    "'- #Face detection using DNN model.py     # legacy DNN-only demo (kept)"
)

# ============================== 10. MODULES ================================
h1("10. Module-by-Module Breakdown")
bullet("Single source of truth for all file paths and thresholds, anchored to the project root "
       "using pathlib. Replaced the old hardcoded C:\\Users\\sambi\\... paths so it now works on any "
       "machine. ensure_dirs() creates runtime folders.", bold_prefix="face/config.py — ")
bullet("Both detectors expose a uniform detect(frame) -> list[Detection], where each Detection is "
       "a clamped box (x1, y1, x2, y2, confidence) with a .crop() helper. draw_detections() overlays "
       "boxes and labels. Haar reports confidence 1.0 (no real score); DNN reports the model's "
       "actual confidence.", bold_prefix="face/detectors.py — ")
bullet("The pure 'signals' core (no camera, no Streamlit imports, so it is reusable and testable): "
       "rgb_gray_views() (channel split + luma grayscale + weights), histogram(), and blob_views() "
       "(unpacks the (1,3,300,300) blob into channels, heatmaps, histograms, and stats).",
       bold_prefix="face/signals.py — ")
bullet("LbphRecognizer and EmbeddingRecognizer, both with fit / save / load / identify. identify() "
       "returns Prediction(name, score, is_known) where LOWER score = MORE confident for both. Each "
       "has an 'Unknown' threshold.", bold_prefix="face/recognizers.py — ")
bullet("capture_from_webcam(), add_images_from_folder(), load_dataset(), and train_all() (rebuilds "
       "and persists both recognizers).", bold_prefix="face/enroll.py — ")
bullet("run_comparison() does a per-person train/test split, evaluates both pipelines, and returns "
       "metrics plus a per-criterion and overall winner.", bold_prefix="face/compare.py — ")
bullet("Shared helpers: bgr_to_rgb() (OpenCV is BGR, Streamlit expects RGB), a Camera "
       "context-manager wrapper around cv2.VideoCapture, and @st.cache_resource singletons so models "
       "load once per session.", bold_prefix="ui/common.py — ")
bullet("Wires the five tabs together; contains no logic itself.", bold_prefix="app.py — ")

# ============================== 11. UI =====================================
h1("11. The 5-Tab User Interface")
table(
    ["Tab", "Purpose"],
    [
        ["Live", "Both pipelines detect + recognize on the same webcam frames in real time; shows names and end-to-end FPS."],
        ["Signal Lab", "RGB -> grayscale signal transformation with channels, luma formula, and histograms."],
        ["BLOB Lab", "How the SSD input blob varies face-to-face: channel heatmaps, stats, histograms."],
        ["Enroll", "Register people (webcam or folder) and train both recognizers."],
        ["Compare", "Accuracy + robustness + FPS metrics and the overall winning pipeline."],
    ],
)

# ============================== 12. INSTALL ================================
h1("12. Installation & Running")
code("# 1. Install dependencies (MUST be the contrib build for cv2.face / LBPH)\n"
     "pip install -r requirements.txt\n\n"
     "# 2. Download the 3 model files into models/\n"
     "python scripts/download_models.py\n\n"
     "# 3. Launch the app\n"
     "streamlit run app.py")
p("Typical usage flow:")
numbered("Open the Enroll tab -> capture yourself via webcam (or import a folder) -> click Train.")
numbered("Go to Live -> both pipelines now label you by name; strangers show 'Unknown'.")
numbered("Explore Signal Lab and BLOB Lab to show the signal internals.")
numbered("Run Compare to see which pipeline wins.")
p("The three model files (auto-downloaded):")
bullet("deploy.prototxt — SSD detector architecture.")
bullet("res10_300x300_ssd_iter_140000.caffemodel — SSD detector weights.")
bullet("openface_nn4.small2.v1.t7 — 128-D face embedder.")

# ============================== 13. DATA FLOW =============================
h1("13. Data Flow (end to end)")
numbered("Webcam frame is captured (BGR image) via cv2.VideoCapture.")
numbered("DETECTION: Haar runs on grayscale; DNN-SSD runs on the 300x300 mean-subtracted blob. "
         "Each returns bounding boxes.")
numbered("CROP: each detected face region is cut out of the frame.")
numbered("RECOGNITION: LBPH predicts a label+distance from the grayscale crop; OpenFace produces a "
         "128-D embedding compared to enrolled centroids by cosine distance.")
numbered("THRESHOLD: if the score is worse than the threshold, the label becomes 'Unknown'.")
numbered("DISPLAY: boxes + names are drawn (after BGR->RGB conversion) and shown in Streamlit.")
numbered("COMPARE (offline): the same detect+recognize steps run over the enrolled dataset to "
         "produce accuracy, robustness, and FPS metrics.")

# ============================== 14. DECISIONS =============================
h1("14. Design Decisions & Why")
bullet("Each detector is paired with a recognizer of matching 'spirit' — Haar with histogram-based "
       "LBPH, deep SSD with deep OpenFace embeddings — so the comparison is fair and each pipeline "
       "is internally consistent.", bold_prefix="Pairing — ")
bullet("LBPH and OpenFace both run through OpenCV (cv2.face and cv2.dnn), avoiding dlib / "
       "face_recognition and their heavy compilation. Lighter install, fewer failure points.",
       bold_prefix="OpenCV-only — ")
bullet("Chosen for a clean multi-tab UI that can show live video, image panels, tables, and charts "
       "with very little code, ideal for a teaching/PBL demo.", bold_prefix="Streamlit — ")
bullet("All paths anchored to the project root via pathlib, replacing the original machine-specific "
       "C:\\Users\\sambi\\... paths so the project is portable.", bold_prefix="Centralised config — ")
bullet("signals.py is kept free of camera/Streamlit imports so the signal logic is pure, testable, "
       "and reusable by multiple tabs.", bold_prefix="Pure signal core — ")

# ============================== 15. GOTCHAS ==============================
h1("15. Key Technical Notes & Gotchas (good for Q&A)")
bullet("cv2.face (LBPH) is absent in plain opencv-python.", bold_prefix="opencv-contrib-python is mandatory — ")
bullet("OpenCV uses BGR; Streamlit's st.image expects RGB. The bgr_to_rgb() helper converts before "
       "display, otherwise colours look wrong (red/blue swapped).", bold_prefix="BGR vs RGB — ")
bullet("for BOTH recognizers, lower score = more confident (LBPH distance and embedding cosine "
       "distance). Above-threshold -> 'Unknown'.", bold_prefix="Score direction — ")
bullet("Streamlit live capture uses cv2.VideoCapture(0) on the server, so it must run LOCALLY on "
       "the machine with the webcam. True in-browser webcam would need streamlit-webrtc (not "
       "currently included).", bold_prefix="Server-side webcam — ")
bullet("after retraining, the cached recognizers must be cleared so the new model loads "
       "(clear_recognizer_cache()).", bold_prefix="Caching — ")
bullet("the two original single-file demos are kept and still run, but the DNN one still has the "
       "old hardcoded C:\\Users\\sambi\\... paths; the new app does not (it uses config.py).",
       bold_prefix="Legacy scripts — ")

# ============================== 16. LIMITATIONS =========================
h1("16. Limitations & Honest Caveats")
bullet("Accuracy is measured on a held-out split of the enrolled faces, not an independent test "
       "set — it is indicative, not a benchmark.")
bullet("Haar struggles with non-frontal faces and poor lighting; expect more misses than the DNN.")
bullet("Recognition quality depends on enrollment quality — few or similar-looking samples reduce "
       "accuracy.")
bullet("The app must run on the same machine as the webcam (server-side capture).")
bullet("No GPU acceleration is configured; the DNN runs on CPU.")

# ============================== 17. FUTURE ==============================
h1("17. Possible Future Improvements")
bullet("Add streamlit-webrtc for true in-browser webcam streaming.")
bullet("Add an independent labelled test set for a real accuracy benchmark.")
bullet("Add face alignment before embedding to boost recognition accuracy.")
bullet("Add liveness / anti-spoofing detection.")
bullet("Export comparison results to CSV / a downloadable report.")
bullet("GPU acceleration for the DNN forward pass.")

# ============================== 18. VERIFICATION =======================
h1("18. Verification Performed")
bullet("Dependencies installed; cv2.face confirmed present (opencv-contrib).")
bullet("All three model files downloaded successfully into models/.")
bullet("Every Python module byte-compiles without error.")
bullet("Detectors, signals, recognizers, and the comparison were smoke-tested end to end.")
bullet("The Streamlit app boots headless and serves successfully (HTTP 200).")

# ============================== 19. GLOSSARY ===========================
h1("19. Glossary")
table(
    ["Term", "Meaning"],
    [
        ["Detection", "Locating faces (bounding boxes) in an image."],
        ["Recognition", "Identifying which known person a detected face belongs to."],
        ["Haar Cascade", "Classic fast face detector using light/dark rectangle features in stages."],
        ["LBPH", "Local Binary Patterns Histograms — a histogram-based face recognizer."],
        ["SSD", "Single Shot Detector — a deep network that detects objects in one pass."],
        ["BLOB", "The preprocessed (resized, mean-subtracted) tensor fed into a DNN."],
        ["Embedding", "A 128-D numeric vector representing a face; similar faces are close together."],
        ["Centroid", "The average embedding of all of one person's enrolled samples."],
        ["Cosine distance", "1 - cosine similarity; small distance means very similar vectors."],
        ["Threshold", "Score cut-off above which a match is rejected as 'Unknown'."],
        ["Luma", "Weighted grayscale brightness: 0.299R + 0.587G + 0.114B."],
        ["BGR / RGB", "OpenCV stores colour as Blue-Green-Red; most display libraries use Red-Green-Blue."],
    ],
)

# ============================== 20. SUMMARY ============================
h1("20. One-Paragraph Summary (for oral explanation)")
p("\"This project compares two complete face pipelines — a classic Haar cascade paired with "
  "histogram-based LBPH recognition, and a deep-learning SSD detector paired with 128-dimensional "
  "OpenFace embeddings. Beyond just detecting faces, it recognizes who they are after a quick "
  "enrollment step. It is wrapped in a five-tab Streamlit interface that also exposes the "
  "signal-processing internals: one tab shows how the RGB camera signal collapses into grayscale "
  "using the luma formula, and another shows how the SSD's input 'blob' — the resized, "
  "mean-subtracted tensor the network actually sees — varies numerically from face to face. "
  "Finally, a comparison tab scores both pipelines on recognition accuracy, detection robustness, "
  "and speed, and declares an overall winner. Everything stays within OpenCV, so it is lightweight "
  "and easy to install.\"", italic=True)

doc.save(str(OUT))
print("WROTE", OUT)
